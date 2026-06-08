#!/usr/bin/env python3
"""Synthetic golden fixtures for the document-processing quality eval.

Builds small, deterministic spend + context documents (csv / xlsx / json / txt /
docx) into ``tests/eval/golden/document_processing/`` together with an
``expected.json`` describing what good processing should produce. The fixtures
are committable; the eval runner regenerates any that are missing so it is
runnable from a clean checkout.

Kept dependency-light: csv/json/txt via stdlib, xlsx via openpyxl, docx via
python-docx (both already used by the ingestion pipeline).
"""
from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any, Dict, List

ROOT = Path(__file__).parent.parent
FIXTURES_DIR = ROOT / "tests" / "eval" / "golden" / "document_processing"
EXPECTED_PATH = FIXTURES_DIR / "expected.json"

# Shared spend rows reused across csv / xlsx / json so totals stay in sync.
_SPEND_ROWS = [
    ("Infosys", "IT managed services Q1", 3_200_000, "IT", "2025-01-15", "INR"),
    ("AWS", "Cloud hosting", 2_100_000, "IT", "2025-02-10", "INR"),
    ("Deloitte", "Advisory engagement", 1_800_000, "Professional Services", "2025-02-20", "INR"),
    ("MakeMyTrip", "Air travel bookings", 950_000, "Travel", "2025-03-05", "INR"),
    ("WeWork", "Office space lease", 1_400_000, "Facilities", "2025-03-12", "INR"),
    ("Staples", "Office supplies", 250_000, "Procurement", "2025-03-18", "INR"),
]
_SPEND_TOTAL = float(sum(r[2] for r in _SPEND_ROWS))
_SPEND_COUNT = len(_SPEND_ROWS)

# A markdown context memo — headings + a table + prose, so the hierarchical
# chunker produces parents, table-row children, and populated heading paths.
_CONTEXT_MEMO_MD = """# Cost Optimization Memo — Aranya Digital Services

## Executive Summary

This memo reviews operating expense across IT, Travel, and Professional Services
for FY2025. Total addressable spend is estimated at INR 18 crore. The single
largest opportunity sits in IT software licensing, where contract renewal timing
creates negotiating leverage.

## Cost Summary by Category

| Category | Annual Spend (INR Lakh) | Addressable % |
| --- | --- | --- |
| IT & Technology | 1,200 | 35 |
| Travel | 800 | 20 |
| Professional Services | 640 | 28 |

## Contract Notes

The Oracle database maintenance contract renews in March 2027 and carries a
12 percent auto-escalation clause. The AWS enterprise agreement is up for
renewal in September 2026. Both contracts should be renegotiated well ahead of
their expiry dates to avoid the escalation.

## Recommendations

Consolidate the three regional travel desks into a single managed travel program
and renegotiate Oracle maintenance from Net-30 to Net-60 payment terms to release
working capital.
"""

_NOTES_TXT = (
    "Procurement working notes. The travel category is fragmented across three "
    "regional desks with no single preferred-supplier agreement. Professional "
    "services spend is dominated by a single advisory firm. IT licensing renewals "
    "are clustered in the first half of the fiscal year.\n"
)


def _build_csv(path: Path) -> None:
    with path.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.writer(fh)
        writer.writerow(["Supplier", "Description", "Amount", "Category", "Invoice Date", "Currency"])
        for supplier, desc, amount, cat, date, ccy in _SPEND_ROWS:
            writer.writerow([supplier, desc, amount, cat, date, ccy])


def _build_malformed_csv(path: Path) -> None:
    # Rows present but NO numeric amount column — amounts are text labels, so the
    # pipeline must surface a zero_spend_warning rather than silently ingesting 0.
    with path.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.writer(fh)
        writer.writerow(["Supplier", "Description", "Line Item"])
        writer.writerow(["Infosys", "IT services", "License renewal"])
        writer.writerow(["AWS", "Cloud", "Compute hours"])
        writer.writerow(["Deloitte", "Advisory", "Strategy review"])


def _build_xlsx(path: Path) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    # Sheet 1: the transaction ledger (should be auto-selected for ingestion).
    ws = wb.active
    ws.title = "Transactions"
    ws.append(["Supplier", "Description", "Amount", "Date"])
    for supplier, desc, amount, _cat, date, _ccy in _SPEND_ROWS:
        ws.append([supplier, desc, amount, date])
    # Sheet 2: a rollup summary (should be skipped as non-ledger).
    ws2 = wb.create_sheet("Summary")
    ws2.append(["Category", "Total"])
    ws2.append(["IT", 5_300_000])
    ws2.append(["Other", 4_400_000])
    # Sheet 3: assumptions (should be skipped).
    ws3 = wb.create_sheet("Assumptions")
    ws3.append(["Parameter", "Value"])
    ws3.append(["Discount rate", 0.12])
    wb.save(path)


def _build_json(path: Path) -> None:
    records: List[Dict[str, Any]] = [
        {
            "supplier": supplier,
            "description": desc,
            "amount": amount,
            "category": cat,
            "date": date,
            "currency": ccy,
        }
        for supplier, desc, amount, cat, date, ccy in _SPEND_ROWS
    ]
    path.write_text(json.dumps(records, indent=2), encoding="utf-8")


def _build_txt(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def _build_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Cost Optimization Memo — Aranya Digital Services", level=1)
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(
        "This memo reviews operating expense across IT, Travel, and Professional "
        "Services for FY2025. Total addressable spend is estimated at INR 18 crore."
    )
    doc.add_heading("Contract Notes", level=2)
    doc.add_paragraph(
        "The Oracle database maintenance contract renews in March 2027 and carries "
        "a 12 percent auto-escalation clause. The AWS enterprise agreement is up for "
        "renewal in September 2026."
    )
    doc.add_heading("Recommendations", level=2)
    doc.add_paragraph(
        "Consolidate the three regional travel desks into a single managed travel "
        "program and renegotiate Oracle maintenance from Net-30 to Net-60 payment terms."
    )
    doc.save(path)


# ---------------------------------------------------------------------------
# Fixture specs + expectations
# ---------------------------------------------------------------------------

def fixture_specs() -> List[Dict[str, Any]]:
    """The golden fixtures and what good processing should produce for each."""
    return [
        {
            "name": "clean_spend.csv",
            "kind": "tabular",
            "format": "csv",
            "expected": {
                "parse_ok": True,
                "line_count": _SPEND_COUNT,
                "total_amount": _SPEND_TOTAL,
                "currency": "INR",
                "schema_roles": {
                    "supplier": "Supplier",
                    "description": "Description",
                    "amount": "Amount",
                    "date": "Invoice Date",
                    "currency": "Currency",
                },
                "zero_spend_warning": False,
            },
        },
        {
            "name": "multi_sheet.xlsx",
            "kind": "tabular",
            "format": "xlsx",
            "expected": {
                "parse_ok": True,
                "line_count": _SPEND_COUNT,
                "total_amount": _SPEND_TOTAL,
                "currency": "INR",
                "schema_roles": {
                    "supplier": "Supplier",
                    "description": "Description",
                    "amount": "Amount",
                    "date": "Date",
                },
                "ledger_sheet": "Transactions",
                "ledger_header_row": 0,
                "zero_spend_warning": False,
            },
        },
        {
            "name": "spend.json",
            "kind": "tabular",
            "format": "json",
            "expected": {
                "parse_ok": True,
                "line_count": _SPEND_COUNT,
                "total_amount": _SPEND_TOTAL,
                "currency": "INR",
                # JSON parse path does not run column schema inference.
                "schema_roles": {},
                "zero_spend_warning": False,
            },
        },
        {
            "name": "malformed.csv",
            "kind": "tabular",
            "format": "csv",
            "expected": {
                "parse_ok": True,
                "zero_spend_warning": True,
                "expect_warnings": True,
            },
        },
        {
            "name": "context_memo.txt",
            "kind": "document",
            "format": "txt",
            "expected": {
                "parse_ok": True,
                "min_chars": 400,
                "has_markdown_headings": True,
                "min_parents": 3,
                "min_children": 4,
                "expect_heading_path": True,
                "expect_table_rows": True,
                "retrieval": [
                    {"query": "IT & Technology annual spend in the cost summary", "expect_contains": "1,200"},
                    {"query": "When does the Oracle maintenance contract renew?", "expect_contains": "March 2027"},
                ],
            },
        },
        {
            "name": "notes.txt",
            "kind": "document",
            "format": "txt",
            "expected": {
                "parse_ok": True,
                "min_chars": 150,
                "has_markdown_headings": False,
                "min_parents": 1,
                "min_children": 1,
                "expect_heading_path": False,
                "expect_table_rows": False,
                "retrieval": [
                    {"query": "How many regional travel desks are there?", "expect_contains": "three regional desks"},
                ],
            },
        },
        {
            "name": "context_memo.docx",
            "kind": "document",
            "format": "docx",
            "expected": {
                "parse_ok": True,
                "min_chars": 200,
                "has_markdown_headings": False,
                "min_parents": 1,
                "min_children": 1,
                "expect_heading_path": False,
                "expect_table_rows": False,
                "retrieval": [
                    {"query": "Oracle contract renewal date", "expect_contains": "March 2027"},
                ],
            },
        },
    ]


def ensure_fixtures() -> List[Dict[str, Any]]:
    """Build any missing fixture files and write expected.json. Returns the specs."""
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    specs = fixture_specs()
    builders = {
        "clean_spend.csv": lambda p: _build_csv(p),
        "malformed.csv": lambda p: _build_malformed_csv(p),
        "multi_sheet.xlsx": lambda p: _build_xlsx(p),
        "spend.json": lambda p: _build_json(p),
        "context_memo.txt": lambda p: _build_txt(p, _CONTEXT_MEMO_MD),
        "notes.txt": lambda p: _build_txt(p, _NOTES_TXT),
        "context_memo.docx": lambda p: _build_docx(p),
    }
    for spec in specs:
        path = FIXTURES_DIR / spec["name"]
        if not path.exists():
            builders[spec["name"]](path)
    EXPECTED_PATH.write_text(
        json.dumps({s["name"]: s for s in specs}, indent=2), encoding="utf-8"
    )
    return specs


if __name__ == "__main__":
    built = ensure_fixtures()
    print(f"Built {len(built)} fixtures in {FIXTURES_DIR}")
