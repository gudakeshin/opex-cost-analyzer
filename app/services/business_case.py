from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any, Dict

from docx import Document

from app.config import OUTPUT_DIR
from app.utils.inr_format import format_money


def _resolve_reporting_currency(analysis: Dict[str, Any]) -> str:
    """Session reporting currency (manifest / analysis payload), default INR for this product."""
    for key in ("reporting_currency", "currency"):
        val = analysis.get(key)
        if val:
            return str(val).upper()
    manifest_raw = analysis.get("manifest")
    manifest = manifest_raw if isinstance(manifest_raw, dict) else {}
    for key in ("reporting_currency", "currency"):
        val = manifest.get(key)
        if val:
            return str(val).upper()
    return "INR"


def _fmt_money(amount: float, currency: str, *, parens: bool = False) -> str:
    text = format_money(float(amount or 0.0), currency)
    return f"({text})" if parens else text


def _top_categories_from_profile(outputs: Dict[str, Any], limit: int = 3) -> list[str]:
    profile = outputs.get("spend-profiler", {})
    cats = profile.get("category_profile", [])
    return [c.get("category_name", c.get("category_id", "")) for c in cats[:limit] if c.get("category_name") or c.get("category_id")]


def _dynamic_risks(outputs: Dict[str, Any]) -> list[str]:
    risks: list[str] = []
    ctx = outputs.get("document-contextualizer", {})
    for constraint in ctx.get("constraints", []):
        risks.append(str(constraint))
    peer = outputs.get("peer-benchmarker", {})
    metadata = peer.get("benchmark_metadata", {}) if isinstance(peer, dict) else {}
    if metadata:
        note = metadata.get("data_quality_note", "")
        if note:
            risks.append(f"Benchmark data quality: {note}")
        score = metadata.get("specificity_score")
        if score is not None and float(score) < 0.7:
            risks.append(f"Benchmark specificity score {score:.2f} — consider acquiring a targeted licensed dataset.")
    checks = outputs.get("data-validator", {}).get("checks", {})
    if checks and not all(bool(v) for v in checks.values()):
        risks.append("Data validation checks indicate potential reliability gaps.")
    if not risks:
        risks.append("No critical model risks detected from current uploaded data.")
    return risks


def _build_financial_projections(outputs: Dict[str, Any], currency: str) -> Dict[str, Any]:
    """Y1/Y2/Y3 table from savings_modeler initiatives."""
    initiatives = outputs.get("savings-modeler", {}).get("initiatives", [])
    gross = [0.0, 0.0, 0.0]
    cta = [0.0, 0.0, 0.0]
    net = [0.0, 0.0, 0.0]
    for init in initiatives:
        for i, k in enumerate(["y1", "y2", "y3"]):
            gross[i] += float(init.get("gross_savings", {}).get(k, 0.0))
            cta[i] += float(init.get("cost_to_achieve", {}).get(k, 0.0))
            net[i] += float(init.get("net_savings", {}).get(k, 0.0))
    cum = [net[0], net[0] + net[1], net[0] + net[1] + net[2]]
    return {
        "headers": ["", "Year 1", "Year 2", "Year 3", "3-Year Total"],
        "rows": [
            ["Gross Savings"] + [_fmt_money(v, currency) for v in gross] + [_fmt_money(sum(gross), currency)],
            ["Cost to Achieve"] + [_fmt_money(v, currency, parens=True) for v in cta] + [_fmt_money(sum(cta), currency, parens=True)],
            ["Net Savings"] + [_fmt_money(v, currency) for v in net] + [_fmt_money(sum(net), currency)],
            ["Cumulative Net"] + [_fmt_money(v, currency) for v in cum] + [""],
        ],
    }


def _build_irr_summary(outputs: Dict[str, Any]) -> list[str]:
    lines_out = []
    for init in outputs.get("savings-modeler", {}).get("initiatives", []):
        irr = init.get("irr_pct")
        if irr is not None:
            lines_out.append(
                f"{init.get('category_name', init.get('category_id', '?'))}: "
                f"IRR = {irr}%, Payback = {init.get('payback_months', '?')} months"
            )
    return lines_out or ["No CTA-based initiatives; IRR not applicable."]


def _build_do_nothing(conf: Dict[str, Any], currency: str) -> str:
    mid = float(conf.get("mid", 0.0))
    return (
        f"Cost of inaction over 3 years: {_fmt_money(mid, currency)} in foregone savings. "
        f"Competitors operating at P25 efficiency hold a structural cost advantage "
        f"of approximately {_fmt_money(mid / 3, currency)} per annum."
    )


def _build_assumption_register_section(outputs: Dict[str, Any]) -> Dict[str, Any]:
    """Pull assumption register data into business-case section."""
    ar = outputs.get("assumption-register", {})
    if not ar:
        return {"available": False, "note": "Run assumption-register skill to populate this section."}
    return {
        "available": True,
        "portfolio_aqs": ar.get("portfolio_aqs"),
        "method": ar.get("method"),
        "initiative_count": len(ar.get("initiative_assumptions", [])),
        "gate2_threshold": 0.65,
        "gate2_status": (
            "PASS" if (ar.get("portfolio_aqs") or 0) >= 0.65 else "BLOCKED — CFO override required"
        ),
        "p10_p50_p90_summary": [
            {
                "initiative": ia.get("initiative_id"),
                "p10": ia.get("p10"),
                "p50": ia.get("p50"),
                "p90": ia.get("p90"),
                "aqs": ia.get("composite_score"),
            }
            for ia in ar.get("initiative_assumptions", [])[:5]
        ],
    }


def _build_rag_section(outputs: Dict[str, Any]) -> Dict[str, Any]:
    """Risk-adjusted growth: scenario floor and macro sensitivity."""
    sm = outputs.get("scenario-modeler", {})
    return {
        "macro_sensitivity": sm.get("macro_sensitivity_rating", "unknown"),
        "downside_floor": sm.get("downside_floor"),
        "downside_pct_of_base": sm.get("downside_floor_pct_of_base"),
        "p10_savings": sm.get("p10_savings"),
        "p90_savings": sm.get("p90_savings"),
    }


# ---------------------------------------------------------------------------
# Layer B — LLM advisory narrative (reuses the reflect-phase synthesis)
# ---------------------------------------------------------------------------

def _llm_advisory_sections(analysis: Dict[str, Any]) -> Any | None:
    """Run the existing LLM advisory synthesis for the business case.

    Returns an ``AdvisorySections`` (or None when the LLM is disabled / low
    quality). Best-effort: any failure degrades to the deterministic sections.
    """
    try:
        from app.opar.models import ObserveContext
        from app.opar.reflect_advisory import generate_llm_advisory_sections
    except Exception:
        return None
    outputs = analysis.get("skill_outputs", {})
    manifest = {
        "company_name": analysis.get("company_name"),
        "industry": analysis.get("industry"),
        "annual_revenue": analysis.get("annual_revenue"),
        "currency": analysis.get("reporting_currency"),
        "engagement_id": analysis.get("engagement_id", "") or "",
    }
    try:
        ctx = ObserveContext(
            user_message=(
                "Produce a detailed CFO business case for the modeled OpEx savings "
                "initiatives, with per-initiative operational and commercial detail."
            ),
            intent_class="business_case",
            wants_executive_narrative=True,
            model_manifest=analysis.get("model_manifest", {}) or {},
            engagement_id=str(analysis.get("engagement_id", "") or ""),
            deep_research_summary=analysis.get("deep_research_summary"),
        )
        advisory, _thinking, _skip = generate_llm_advisory_sections(
            ctx, manifest, outputs, category_focused=True
        )
        return advisory
    except Exception:
        return None


def _match_advisory_lever(advisory_levers: list, initiative: Dict[str, Any]) -> Any | None:
    """Fuzzy-match an LLM business_lever to a modeled initiative by name/category."""
    name = str(initiative.get("lever_name") or initiative.get("lever") or "")
    cat = str(initiative.get("category_name") or "").lower()
    tokens = {t for t in name.lower().replace("_", " ").split() if len(t) > 2}
    best, best_score = None, 0
    for lev in advisory_levers:
        ln = str(getattr(lev, "lever_name", "") or "").lower()
        score = len(tokens & {t for t in ln.replace("_", " ").split() if len(t) > 2})
        if cat and cat in ln:
            score += 1
        if score > best_score:
            best, best_score = lev, score
    return best if best_score >= 1 else None


def build_initiative_details(outputs: Dict[str, Any], advisory: Any | None = None) -> list[Dict[str, Any]]:
    """Per-initiative business detail: Layer-A fields + value-bridge financials + LLM sharpening.

    This is the single source the business-case document and the pipeline
    persistence both read, so depth and financials stay consistent.
    """
    savings = outputs.get("savings-modeler", {})
    initiatives = savings.get("initiatives", []) if isinstance(savings, dict) else []
    bridge = outputs.get("value-bridge-calculator", {})
    matrix_by_key = {
        (r.get("category_id"), r.get("lever")): r
        for r in bridge.get("value_matrix", [])
        if isinstance(r, dict)
    }
    advisory_levers = list(getattr(advisory, "business_levers", []) or []) if advisory else []

    details: list[Dict[str, Any]] = []
    for init in initiatives:
        if not isinstance(init, dict):
            continue
        cid, lever = init.get("category_id"), init.get("lever")
        vm = matrix_by_key.get((cid, lever), {})
        gross = init.get("gross_savings", {}) or {}
        net = init.get("net_savings", {}) or {}

        rationale = init.get("business_rationale") or init.get("root_cause") or ""
        evidence: list[str] = []
        matched = _match_advisory_lever(advisory_levers, init)
        if matched is not None:
            sharpened = " ".join(
                s for s in (getattr(matched, "what_changes", ""), getattr(matched, "why_it_works", "")) if s
            ).strip()
            if sharpened:
                rationale = sharpened
            evidence = list(getattr(matched, "evidence", []) or [])

        details.append({
            "category_id": cid,
            "category_name": init.get("category_name", cid),
            "lever": lever,
            "lever_name": init.get("lever_name", lever),
            "root_cause": init.get("root_cause"),
            "business_rationale": rationale,
            "evidence": evidence,
            "owner": init.get("owner", {}),
            "owner_role": init.get("owner_role"),
            "business_sponsor": init.get("business_sponsor"),
            "affected_vendors": init.get("affected_vendors", []),
            "contract_levers": init.get("contract_levers", []),
            "risks": init.get("risks", []),
            "kpis": init.get("kpis", []),
            "change_management": init.get("change_management", {}),
            "phasing_narrative": init.get("phasing_narrative"),
            "execution_playbook": init.get("execution_playbook", []),
            "savings_type": init.get("savings_type", "run_rate"),
            "annualized_run_rate_savings": init.get("annualized_run_rate_savings", 0.0),
            # Financials (per-year gross + 3yr + bands) for persistence + display.
            "gross_savings_y1": gross.get("y1", 0.0),
            "gross_savings_y2": gross.get("y2", 0.0),
            "gross_savings_y3": gross.get("y3", 0.0),
            "gross_3yr": gross.get("total_3yr", vm.get("gross_3yr", 0.0)),
            "cost_to_achieve_3yr": (init.get("cost_to_achieve", {}) or {}).get("total_3yr", vm.get("cost_to_achieve_3yr", 0.0)),
            "net_npv": net.get("npv_10pct", vm.get("net_npv", 0.0)),
            "p50_savings": vm.get("deduped_mid_savings"),
            "payback_months": init.get("payback_months", vm.get("payback_months", 0)),
            "irr_pct": init.get("irr_pct"),
            "ebitda_bps": (init.get("ebitda_impact", {}) or {}).get("ebitda_bps"),
            "confidence": init.get("confidence", "medium"),
        })
    return details


def _advisory_business_levers(advisory: Any | None) -> list[Dict[str, Any]]:
    levers: list[Dict[str, Any]] = []
    for lev in getattr(advisory, "business_levers", []) or []:
        levers.append({
            "lever_name": getattr(lev, "lever_name", ""),
            "what_changes": getattr(lev, "what_changes", ""),
            "why_it_works": getattr(lev, "why_it_works", ""),
            "evidence": list(getattr(lev, "evidence", []) or []),
        })
    return levers


def _priority_actions(advisory: Any | None, details: list[Dict[str, Any]], currency: str) -> list[Dict[str, Any]]:
    """30/60/90 actions from the LLM, or a deterministic horizon-led fallback."""
    items = getattr(advisory, "priority_actions_30_60_90", []) or []
    if items:
        return [
            {
                "timeline": getattr(a, "timeline", ""),
                "action": getattr(a, "action", ""),
                "expected_impact": getattr(a, "expected_impact", ""),
            }
            for a in items
        ]
    out: list[Dict[str, Any]] = []
    quick = [d for d in details if (d.get("payback_months") or 99) <= 6][:2]
    for d in quick:
        out.append({
            "timeline": "30",
            "action": f"Mobilize {d['lever_name']} in {d['category_name']} — confirm owner ({d.get('owner_role') or 'owner'}) and baseline.",
            "expected_impact": d.get("phasing_narrative") or "Quick-win realization",
        })
    structural = [d for d in details if d not in quick][:2]
    for d in structural:
        out.append({
            "timeline": "60",
            "action": f"Negotiate / execute {d['lever_name']} in {d['category_name']}.",
            "expected_impact": f"~{_fmt_money(float(d.get('annualized_run_rate_savings', 0) or 0), currency)} run-rate",
        })
    out.append({
        "timeline": "90",
        "action": "Institutionalize governance and monthly pipeline tracking across all committed initiatives.",
        "expected_impact": "Sustained realization and bounce-back control",
    })
    return out


def build_business_case(analysis: Dict[str, Any], template: str = "detailed_proposal") -> Dict[str, Any]:
    outputs = analysis.get("skill_outputs", {})
    bridge = outputs.get("value-bridge-calculator", {})
    conf = bridge.get("confidence_bands", {})
    currency = _resolve_reporting_currency(analysis)
    company_name = analysis.get("company_name") or "Client"
    top_categories = _top_categories_from_profile(outputs)
    top_text = ", ".join(top_categories) if top_categories else "top identified categories"

    advisory = _llm_advisory_sections(analysis)
    details = build_initiative_details(outputs, advisory)

    # Executive summary: LLM takeaway when available, else the deterministic band line.
    deterministic_summary = (
        f"{company_name} has an estimated savings opportunity of "
        f"{_fmt_money(float(conf.get('mid', 0) or 0), currency)} (mid-case), ranging from "
        f"{_fmt_money(float(conf.get('low', 0) or 0), currency)} "
        f"to {_fmt_money(float(conf.get('high', 0) or 0), currency)}."
    )
    takeaway = (getattr(advisory, "executive_takeaway", "") or "").strip()
    executive_summary = f"{takeaway}\n\n{deterministic_summary}" if takeaway else deterministic_summary

    strategic_context = (getattr(advisory, "category_focus_section", "") or "").strip() or (
        f"Current-state baseline built from uploaded spend and contextual documents, "
        f"with focus on {top_text}. The portfolio below details the operational and "
        f"commercial change behind each modeled initiative — affected vendors, accountable "
        f"owner, execution risks and the KPIs that prove realization."
    )

    sections = {
        "executive_summary": executive_summary,
        "strategic_context": strategic_context,
        "current_state": f"Current-state baseline built from uploaded spend and contextual documents, with focus on {top_text}.",
        "financial_projections": _build_financial_projections(outputs, currency),
        "irr_summary": _build_irr_summary(outputs),
        "business_levers": _advisory_business_levers(advisory),
        "initiative_details": details,
        "priority_actions_30_60_90": _priority_actions(advisory, details, currency),
        "quick_wins": list(getattr(advisory, "quick_wins_from_data", []) or []),
        "executive_callouts": list(getattr(advisory, "executive_callouts", []) or []),
        "do_nothing_comparison": _build_do_nothing(conf, currency),
        "savings_opportunity": bridge.get("value_matrix", []),
        "implementation_approach": [
            f"Wave 1: Execute initiatives in {top_text}.",
            "Wave 2: Extend controls to remaining categories with measurable variance.",
            "Wave 3: Institutionalize governance and monthly tracking in pipeline.",
        ],
        "risk_assessment": _dynamic_risks(outputs),
        "assumption_register": _build_assumption_register_section(outputs),
        "rag_factors": _build_rag_section(outputs),
    }
    return {
        "template": template,
        "generated_on": str(date.today()),
        "reporting_currency": currency,
        "sections": sections,
    }


def _render_table(doc: Document, table_data: Dict[str, Any]) -> None:
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    if not headers or not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = str(h)
    for i, row_data in enumerate(rows):
        for j, cell_val in enumerate(row_data[: len(headers)]):
            table.rows[i + 1].cells[j].text = str(cell_val)
    doc.add_paragraph("")


def _money(v: Any, currency: str = "INR") -> str:
    try:
        return _fmt_money(float(v), currency)
    except (TypeError, ValueError):
        return str(v)


def _render_initiative_details_docx(doc: Document, details: list, currency: str = "INR") -> None:
    if not details:
        doc.add_paragraph("No modeled initiatives available.")
        return
    for d in details:
        if not isinstance(d, dict):
            continue
        title = f"{d.get('lever_name') or d.get('lever') or 'Initiative'} — {d.get('category_name') or d.get('category_id') or ''}"
        doc.add_heading(title.strip(" —"), level=2)
        if d.get("business_rationale"):
            doc.add_paragraph(str(d["business_rationale"]))
        owner_bits = [b for b in (d.get("owner_role"), d.get("business_sponsor")) if b]
        if owner_bits:
            doc.add_paragraph(f"Owner / accountability: {owner_bits[0]} (sponsor: {owner_bits[-1]})." if len(owner_bits) > 1 else f"Owner / accountability: {owner_bits[0]}.")
        fin = (
            f"3-yr gross {_money(d.get('gross_3yr'), currency)} · net NPV {_money(d.get('net_npv'), currency)} · "
            f"payback {d.get('payback_months', '?')} mo"
            + (f" · IRR {d.get('irr_pct')}%" if d.get("irr_pct") is not None else "")
            + (f" · {d.get('ebitda_bps')} bps EBITDA" if d.get("ebitda_bps") is not None else "")
        )
        doc.add_paragraph(fin)
        if d.get("affected_vendors"):
            doc.add_paragraph("Affected vendors:")
            for v in d["affected_vendors"]:
                doc.add_paragraph(
                    f"{v.get('supplier')} — {_money(v.get('spend'), currency)} ({v.get('share_of_category_pct')}% of category)",
                    style="List Bullet",
                )
        if d.get("contract_levers"):
            doc.add_paragraph("Contract & commercial levers:")
            for c in d["contract_levers"]:
                doc.add_paragraph(str(c), style="List Bullet")
        if d.get("risks"):
            doc.add_paragraph("Risks & mitigations:")
            for r in d["risks"]:
                doc.add_paragraph(
                    f"[{r.get('severity', 'med')}] {r.get('risk')} → {r.get('mitigation')}",
                    style="List Bullet",
                )
        if d.get("kpis"):
            doc.add_paragraph("KPIs:")
            for k in d["kpis"]:
                doc.add_paragraph(f"{k.get('metric')} ({k.get('cadence')})", style="List Bullet")
        cm = d.get("change_management") or {}
        if cm.get("stakeholders"):
            doc.add_paragraph(
                f"Change management — stakeholders: {', '.join(cm.get('stakeholders', []))}; "
                f"cadence: {cm.get('comms_cadence', 'n/a')}."
            )
        if d.get("phasing_narrative"):
            doc.add_paragraph(f"Phasing: {d['phasing_narrative']}")
        if d.get("evidence"):
            doc.add_paragraph("Evidence:")
            for e in d["evidence"]:
                doc.add_paragraph(str(e), style="List Bullet")


def _render_business_levers_docx(doc: Document, levers: list) -> None:
    for lev in levers:
        if not isinstance(lev, dict):
            continue
        doc.add_heading(str(lev.get("lever_name", "Lever")), level=2)
        if lev.get("what_changes"):
            doc.add_paragraph(f"What changes: {lev['what_changes']}")
        if lev.get("why_it_works"):
            doc.add_paragraph(f"Why it works: {lev['why_it_works']}")
        for e in lev.get("evidence", []) or []:
            doc.add_paragraph(str(e), style="List Bullet")


def _render_priority_actions_docx(doc: Document, actions: list) -> None:
    for a in actions:
        if isinstance(a, dict):
            doc.add_paragraph(
                f"Day {a.get('timeline')}: {a.get('action')} — {a.get('expected_impact')}",
                style="List Bullet",
            )
        else:
            doc.add_paragraph(str(a), style="List Bullet")


# Keys whose value is a list-of-dicts needing a structured renderer.
_STRUCTURED_DOCX: Dict[str, Any] = {
    "initiative_details": _render_initiative_details_docx,
    "business_levers": _render_business_levers_docx,
    "priority_actions_30_60_90": _render_priority_actions_docx,
}


def export_docx(business_case: Dict[str, Any], filename: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    currency = str(business_case.get("reporting_currency") or "INR")
    doc = Document()
    doc.add_heading("OpEx Business Case", level=0)
    doc.add_paragraph(f"Generated: {business_case.get('generated_on')}")
    for key, value in business_case.get("sections", {}).items():
        doc.add_heading(key.replace("_", " ").title(), level=1)
        renderer = _STRUCTURED_DOCX.get(key)
        if renderer is _render_initiative_details_docx and isinstance(value, list):
            renderer(doc, value, currency)
        elif renderer is not None and isinstance(value, list):
            renderer(doc, value)
        elif isinstance(value, str):
            doc.add_paragraph(value)
        elif isinstance(value, dict) and "headers" in value and "rows" in value:
            _render_table(doc, value)
        elif isinstance(value, dict):
            for k, v in value.items():
                doc.add_paragraph(f"{str(k).replace('_', ' ').title()}: {v}", style="List Bullet")
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    doc.add_paragraph(_compact_row(item, currency), style="List Bullet")
                else:
                    doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph(str(value))
    doc.save(path)
    return path


def _compact_row(item: Dict[str, Any], currency: str = "INR") -> str:
    """One-line summary of a value-matrix style row for bullet rendering."""
    name = item.get("lever_name") or item.get("lever") or item.get("category_name") or "item"
    cat = item.get("category_name") or item.get("category_id")
    bits = [str(name)]
    if cat and cat != name:
        bits.append(f"({cat})")
    if item.get("net_npv") is not None:
        bits.append(f"— net NPV {_money(item.get('net_npv'), currency)}")
    if item.get("deduped_mid_savings") is not None:
        bits.append(f"— mid {_money(item.get('deduped_mid_savings'), currency)}")
    if item.get("payback_months"):
        bits.append(f"payback {item.get('payback_months')} mo")
    return " ".join(bits)


def _render_initiative_details_text(d: Dict[str, Any], currency: str = "INR") -> list:
    lines = [
        f"  • {d.get('lever_name') or d.get('lever')} — {d.get('category_name') or d.get('category_id')}",
    ]
    if d.get("business_rationale"):
        lines.append(f"      Rationale: {d['business_rationale']}")
    if d.get("owner_role"):
        lines.append(f"      Owner: {d['owner_role']} (sponsor: {d.get('business_sponsor', 'CFO')})")
    lines.append(
        f"      Financials: 3-yr gross {_money(d.get('gross_3yr'), currency)}, net NPV {_money(d.get('net_npv'), currency)}, "
        f"payback {d.get('payback_months', '?')} mo"
    )
    for v in d.get("affected_vendors", []) or []:
        lines.append(f"      Vendor: {v.get('supplier')} ({v.get('share_of_category_pct')}% of category)")
    for c in d.get("contract_levers", []) or []:
        lines.append(f"      Lever: {c}")
    for r in d.get("risks", []) or []:
        lines.append(f"      Risk [{r.get('severity', 'med')}]: {r.get('risk')} → {r.get('mitigation')}")
    for k in d.get("kpis", []) or []:
        lines.append(f"      KPI: {k.get('metric')} ({k.get('cadence')})")
    cm = d.get("change_management") or {}
    if cm.get("stakeholders"):
        lines.append(f"      Change mgmt: {', '.join(cm.get('stakeholders', []))}; {cm.get('comms_cadence', '')}")
    if d.get("phasing_narrative"):
        lines.append(f"      Phasing: {d['phasing_narrative']}")
    return lines


def export_pdf_like_text(business_case: Dict[str, Any], filename: str) -> Path:
    """
    Lightweight fallback export that produces a printable report text file.
    This keeps the endpoint stable without requiring external PDF binaries.
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    currency = str(business_case.get("reporting_currency") or "INR")
    lines = [f"OpEx Business Case - {business_case.get('generated_on')}"]
    for key, value in business_case.get("sections", {}).items():
        lines.append("")
        lines.append(key.replace("_", " ").upper())
        lines.append("-" * 40)
        if key == "initiative_details" and isinstance(value, list):
            for d in value:
                if isinstance(d, dict):
                    lines.extend(_render_initiative_details_text(d, currency))
        elif key == "business_levers" and isinstance(value, list):
            for lev in value:
                if isinstance(lev, dict):
                    lines.append(f"  • {lev.get('lever_name')}")
                    if lev.get("what_changes"):
                        lines.append(f"      What changes: {lev['what_changes']}")
                    if lev.get("why_it_works"):
                        lines.append(f"      Why it works: {lev['why_it_works']}")
        elif key == "priority_actions_30_60_90" and isinstance(value, list):
            for a in value:
                if isinstance(a, dict):
                    lines.append(f"  • Day {a.get('timeline')}: {a.get('action')} — {a.get('expected_impact')}")
        elif isinstance(value, dict) and "headers" in value and "rows" in value:
            headers = value.get("headers", [])
            lines.append(" | ".join(str(h) for h in headers))
            lines.append("-" * 60)
            for row in value.get("rows", []):
                lines.append(" | ".join(str(c) for c in row))
        elif isinstance(value, dict):
            for k, v in value.items():
                lines.append(f"- {str(k).replace('_', ' ').title()}: {v}")
        elif isinstance(value, list):
            for item in value:
                lines.append(f"- {_compact_row(item, currency) if isinstance(item, dict) else item}")
        else:
            lines.append(str(value))
    path.write_text("\n".join(lines), encoding="utf-8")
    return path
