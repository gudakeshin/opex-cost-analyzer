"""
Monthly Operating Review (MOR) Pack Generator.

Produces a structured MOR report from pipeline actuals and BvA outputs.
Output: JSON data dict + DOCX.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document

from app.config import OUTPUT_DIR


def _fmt_cr(v: float) -> str:
    return f"₹{v / 1e7:.1f} Cr" if v else "₹0 Cr"


def build_mor_pack(
    pipeline_summary: Dict[str, Any],
    bva_outputs: Dict[str, Any],
    *,
    company_name: str = "Client",
    review_month: Optional[str] = None,
    engagement_week: int = 1,
) -> Dict[str, Any]:
    """
    Build a monthly operating review pack from pipeline and BvA data.

    pipeline_summary: output of app.services.pipeline.get_pipeline_summary()
    bva_outputs: output of engine.bva_analyzer()
    """
    month = review_month or date.today().strftime("%B %Y")

    # Pull key metrics from pipeline
    committed = float(pipeline_summary.get("committed_savings", 0))
    identified = float(pipeline_summary.get("identified_savings", 0))
    run_rate = float(pipeline_summary.get("run_rate_committed_savings", 0))
    at_risk = float(pipeline_summary.get("at_risk_savings", 0))
    initiatives_total = int(pipeline_summary.get("total_initiatives", 0))
    initiatives_on_track = int(pipeline_summary.get("on_track", 0))
    initiatives_delayed = int(pipeline_summary.get("delayed", 0))

    # BvA KPIs
    bva_summary = bva_outputs.get("summary", {}) if isinstance(bva_outputs, dict) else {}
    total_variance = float(bva_summary.get("total_price_variance", 0) or 0)
    top_variances = bva_outputs.get("category_variances", [])[:5] if isinstance(bva_outputs, dict) else []

    sections: Dict[str, Any] = {
        "header": {
            "company": company_name,
            "review_month": month,
            "engagement_week": engagement_week,
            "generated_on": str(date.today()),
        },
        "pipeline_kpis": {
            "committed_savings_cr": round(committed / 1e7, 1),
            "identified_savings_cr": round(identified / 1e7, 1),
            "run_rate_annualised_cr": round(run_rate / 1e7, 1),
            "at_risk_cr": round(at_risk / 1e7, 1),
            "total_initiatives": initiatives_total,
            "on_track": initiatives_on_track,
            "delayed": initiatives_delayed,
            "delivery_rate_pct": round(
                100 * initiatives_on_track / max(initiatives_total, 1), 1
            ),
        },
        "bva_highlights": {
            "total_variance_cr": round(total_variance / 1e7, 1),
            "top_variances": [
                {
                    "category": v.get("category_name") or v.get("category_id", "?"),
                    "variance_cr": round(float(v.get("total_variance", 0)) / 1e7, 1),
                    "flag": v.get("flag", ""),
                }
                for v in top_variances
            ],
        },
        "action_items": _generate_action_items(pipeline_summary, bva_outputs),
        "next_gate": f"Week {min(12, engagement_week + 1)} — {'Gate ' + str(min(4, (engagement_week // 3) + 1))}",
    }

    return {
        "type": "mor_pack",
        "generated_on": str(date.today()),
        "sections": sections,
    }


def _generate_action_items(pipeline: Dict, bva: Dict) -> List[str]:
    items = []
    delayed = int(pipeline.get("delayed", 0))
    if delayed > 0:
        items.append(f"Review {delayed} delayed initiative(s) with initiative owners; confirm revised timelines.")
    at_risk = float(pipeline.get("at_risk_savings", 0))
    if at_risk > 0:
        items.append(f"At-risk savings of {_fmt_cr(at_risk)} — validate forecast-to-complete with business unit leads.")
    if isinstance(bva, dict):
        for cat in bva.get("category_variances", [])[:3]:
            flag = cat.get("flag", "")
            if "UNFAV" in str(flag).upper():
                items.append(
                    f"Unfavourable variance in {cat.get('category_name', cat.get('category_id', '?'))} "
                    f"— investigate root cause and initiate corrective action."
                )
    if not items:
        items.append("No critical action items identified this month; maintain current cadence.")
    return items


def export_mor_docx(mor: Dict[str, Any], filename: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc = Document()

    hdr = mor["sections"]["header"]
    doc.add_heading(f"Monthly Operating Review — {hdr['company']}", level=0)
    doc.add_paragraph(
        f"{hdr['review_month']}  |  Engagement Week {hdr['engagement_week']}  |  Generated {hdr['generated_on']}"
    )

    doc.add_heading("Pipeline KPIs", level=1)
    kpis = mor["sections"]["pipeline_kpis"]
    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    hdrs = ["Committed (₹ Cr)", "Identified (₹ Cr)", "Run-Rate (₹ Cr)", "At-Risk (₹ Cr)", "Delivery Rate %"]
    vals = [
        str(kpis["committed_savings_cr"]),
        str(kpis["identified_savings_cr"]),
        str(kpis["run_rate_annualised_cr"]),
        str(kpis["at_risk_cr"]),
        str(kpis["delivery_rate_pct"]),
    ]
    for j, h in enumerate(hdrs):
        tbl.rows[0].cells[j].text = h
    for j, v in enumerate(vals):
        tbl.rows[1].cells[j].text = v
    doc.add_paragraph("")

    doc.add_heading("Budget vs. Actuals Highlights", level=1)
    bva = mor["sections"]["bva_highlights"]
    doc.add_paragraph(f"Total cost variance: ₹{bva['total_variance_cr']:.1f} Cr")
    for v in bva.get("top_variances", []):
        doc.add_paragraph(
            f"  • {v['category']}: ₹{v['variance_cr']:.1f} Cr  [{v['flag']}]"
        )

    doc.add_heading("Action Items", level=1)
    for item in mor["sections"].get("action_items", []):
        doc.add_paragraph(f"  ☐ {item}")

    doc.add_paragraph(f"\nNext gate: {mor['sections']['next_gate']}")
    doc.save(path)
    return path
