"""
CFO Brief Generator — 1-page executive summary (DOCX + metadata JSON).

Pulls from OPAR skill outputs: savings-modeler, value-bridge-calculator,
assumption-register, value-to-shareholder-bridge, scenario-modeler,
brsr-cobenefit-calculator.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt

from app.config import OUTPUT_DIR
from app.utils.inr_format import fmt_inr_cr


def _fmt_cr(value: float) -> str:
    return fmt_inr_cr(value, style="cfo")


def _safe_float(outputs: Dict, *keys: str, default: float = 0.0) -> float:
    d: Any = outputs
    for k in keys:
        if not isinstance(d, dict):
            return default
        d = d.get(k, {})
    try:
        return float(d) if d else default
    except (TypeError, ValueError):
        return default


def _top_initiatives(outputs: Dict) -> List[Dict]:
    return outputs.get("savings-modeler", {}).get("initiatives", [])[:5]


def _shareholder_metrics(outputs: Dict) -> Dict:
    bridge = outputs.get("value-to-shareholder-bridge", {})
    return bridge.get("metrics", {})


def _scenario_floor(outputs: Dict) -> Optional[float]:
    sm = outputs.get("scenario-modeler", {})
    return sm.get("downside_floor")


def build_cfo_brief(
    analysis: Dict[str, Any],
    *,
    engagement_week: int = 1,
    pack_id: Optional[str] = None,
    company_name: str = "Client",
) -> Dict[str, Any]:
    """
    Build a structured CFO brief dict from OPAR analysis outputs.
    Returns a dict with sections suitable for DOCX export or JSON.
    """
    outputs = analysis.get("skill_outputs", {})
    bridge = outputs.get("value-bridge-calculator", {})
    conf = bridge.get("confidence_bands", {})

    mid_savings = float(conf.get("mid", 0))
    low_savings = float(conf.get("low", 0))
    high_savings = float(conf.get("high", 0))

    initiatives = _top_initiatives(outputs)
    shareholder = _shareholder_metrics(outputs)
    scenario_floor = _scenario_floor(outputs)

    # AssumptionQualityScore — pulled from assumption-register skill output
    aqs = outputs.get("assumption-register", {}).get("portfolio_aqs", None)
    aqs_text = f"{aqs:.2f}" if aqs else "N/A"

    # Regulatory flags from reg_watcher (surfaced at Reflect gate)
    reg_flags = analysis.get("regulatory_events", [])

    # BRSR headline numbers
    brsr = outputs.get("brsr-cobenefit-calculator", {})
    totals = brsr.get("portfolio_totals", {})
    scope2 = totals.get("delta_scope2_tco2e", 0.0)
    scope3 = totals.get("delta_scope3_tco2e", 0.0)

    sections: Dict[str, Any] = {
        "headline": {
            "company": company_name,
            "engagement_week": engagement_week,
            "pack": pack_id or "generic",
            "as_of": str(date.today()),
            "savings_mid_cr": round(mid_savings / 1e7, 1),
            "savings_low_cr": round(low_savings / 1e7, 1),
            "savings_high_cr": round(high_savings / 1e7, 1),
            "savings_display": f"₹{mid_savings / 1e7:.0f} Cr" if mid_savings else "TBD",
            "range_display": f"₹{low_savings / 1e7:.0f}–{high_savings / 1e7:.0f} Cr",
        },
        "shareholder_bridge": {
            "ebitda_bps": shareholder.get("delta_ebitda_bps"),
            "roce_pp": shareholder.get("delta_roce_pp"),
            "eps_inr": shareholder.get("delta_eps_inr"),
            "equity_value_cr": shareholder.get("delta_equity_value_cr"),
        },
        "top_initiatives": [
            {
                "name": i.get("category_name") or i.get("category_id", "?"),
                "p50_cr": round(float(i.get("p50") or 0) / 1e7, 1),
                "irr_pct": i.get("irr_pct"),
                "payback_months": i.get("payback_months"),
            }
            for i in initiatives
        ],
        "scenario_floor_cr": round(scenario_floor / 1e7, 1) if scenario_floor else None,
        "assumption_quality_score": aqs_text,
        "brsr_cobenefits": {"delta_scope2_tco2e": scope2, "delta_scope3_tco2e": scope3},
        "regulatory_alerts": [r.get("title", r.get("event_id", "?")) for r in reg_flags[:3]],
        "next_decision_gate": f"Gate-{min(4, (engagement_week // 3) + 1)} (Week {min(12, engagement_week + 3 - (engagement_week % 3) or 3)})",
    }

    return {
        "type": "cfo_brief",
        "template": "one_page",
        "generated_on": str(date.today()),
        "sections": sections,
    }


def export_cfo_brief_docx(brief: Dict[str, Any], filename: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc = Document()

    headline = brief["sections"]["headline"]
    doc.add_heading(f"OpEx CFO Brief — {headline['company']}", level=0)
    doc.add_paragraph(
        f"As of {headline['as_of']}  |  Engagement Week {headline['engagement_week']}  |  Pack: {headline['pack']}"
    )

    doc.add_heading("Savings Opportunity", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"₹{headline['savings_mid_cr']:.0f} Cr  ")
    run.bold = True
    run.font.size = Pt(18)
    p.add_run(f"(Range: {headline['range_display']})")

    bridge = brief["sections"]["shareholder_bridge"]
    doc.add_heading("Shareholder Value Bridge", level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ["ΔEBITDA (bps)", "ΔROCE (pp)", "ΔEPS (₹)", "ΔEquity Value (₹ Cr)"]
    values = [
        str(bridge.get("ebitda_bps") or "—"),
        str(bridge.get("roce_pp") or "—"),
        str(bridge.get("eps_inr") or "—"),
        str(bridge.get("equity_value_cr") or "—"),
    ]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for j, v in enumerate(values):
        table.rows[1].cells[j].text = v
    doc.add_paragraph("")

    doc.add_heading("Top Initiatives", level=1)
    inits = brief["sections"]["top_initiatives"]
    if inits:
        tbl = doc.add_table(rows=1 + len(inits), cols=4)
        tbl.style = "Table Grid"
        for j, h in enumerate(["Initiative", "P50 Savings (₹ Cr)", "IRR %", "Payback (mo)"]):
            tbl.rows[0].cells[j].text = h
        for i, row in enumerate(inits):
            tbl.rows[i + 1].cells[0].text = row["name"]
            tbl.rows[i + 1].cells[1].text = str(row["p50_cr"])
            tbl.rows[i + 1].cells[2].text = str(row.get("irr_pct") or "—")
            tbl.rows[i + 1].cells[3].text = str(row.get("payback_months") or "—")
        doc.add_paragraph("")

    doc.add_heading("Risk & Quality Signals", level=1)
    floor = brief["sections"].get("scenario_floor_cr")
    doc.add_paragraph(f"Downside floor (worst scenario): ₹{floor} Cr" if floor else "Downside floor: not computed")
    doc.add_paragraph(f"Assumption Quality Score: {brief['sections']['assumption_quality_score']}")

    reg = brief["sections"].get("regulatory_alerts", [])
    if reg:
        doc.add_paragraph(f"Regulatory alerts: {'; '.join(reg)}")

    brsr = brief["sections"]["brsr_cobenefits"]
    doc.add_heading("BRSR Co-Benefits", level=1)
    doc.add_paragraph(f"ΔScope-2: {brsr['delta_scope2_tco2e']:.1f} tCO₂e  |  ΔScope-3: {brsr['delta_scope3_tco2e']:.1f} tCO₂e")

    doc.add_paragraph(f"\nNext decision gate: {brief['sections']['next_decision_gate']}")

    doc.save(path)
    return path
